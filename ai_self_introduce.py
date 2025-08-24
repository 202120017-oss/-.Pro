# ai_self_introduce.py이 (Pro_V3)
# -------------------------------------------------------------
# 📝 자소서 어렵지 않아용 pro_V3 — 다국어 + 스타일 비교 + 다듬기/분석 + PDF/DOCX
# - 전문 서비스형 UI (상단 헤더 + ⚙️ 팝오버 설정)
# - 한국어 중심 UI, 결과 영문 변환 지원
# - 모델: OpenAI / Gemini (사용자 API Key)
# - 기능:
#   1) 스타일 비교(2–3개 톤 동시 생성)
#   2) 1‑클릭 PDF/DOCX 내보내기(한글 폰트 업로드 지원)
#   3) 다듬기/재톤적용, JD 키워드·겹침 간단 분석
#   4) 인적사항(이름/나이/성별/주소) 입력 포함
#   5) 모든 입력란에 반투명 예시(placeholder) 제공 — 타이핑 시 자동 사라짐
# -------------------------------------------------------------
# 실행
#   pip install -U streamlit openai google-generativeai reportlab python-docx
#   streamlit run app_multilingual_styles_pdf.py

from __future__ import annotations
import os
import io
import re
from typing import List, Dict, Optional

import streamlit as st

# ========== Optional exports ==========
try:
    from docx import Document  # pip install python-docx
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

# PDF (ReportLab)
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.units import mm
    HAS_PDF = True
except Exception:
    HAS_PDF = False

# =============================
# Providers (OpenAI / Gemini)
# =============================
class OpenAIProvider:
    def __init__(self, api_key: Optional[str]):
        self.api_key = api_key or os.getenv("OPENAI_API_KEY", "")
        try:
            from openai import OpenAI
            self.client = OpenAI(api_key=self.api_key)
        except Exception as e:
            self.client = None
            st.error(f"OpenAI 클라이언트 초기화 실패: {e}")

    def generate(self, messages: List[Dict], model: str, temperature: float, max_tokens: int = 1200) -> str:
        if not self.client:
            raise RuntimeError("OpenAI 클라이언트가 준비되지 않았습니다. API 키를 확인하세요.")
        try:
            resp = self.client.chat.completions.create(
                model=model,
                temperature=temperature,
                max_tokens=max_tokens,
                messages=messages,
            )
            return (resp.choices[0].message.content or "").strip()
        except Exception as e:
            raise RuntimeError(f"OpenAI 호출 실패: {e}")

class GeminiProvider:
    def __init__(self, api_key: Optional[str]):
        self.api_key = api_key or os.getenv("GEMINI_API_KEY", "")
        try:
            import google.generativeai as genai
            genai.configure(api_key=self.api_key)
            self.genai = genai
        except Exception as e:
            self.genai = None
            st.error(f"Gemini 클라이언트 초기화 실패: {e}")

    def generate(self, messages: List[Dict], model: str, temperature: float, max_tokens: int = 1200) -> str:
        if not self.genai:
            raise RuntimeError("Gemini 클라이언트가 준비되지 않았습니다. API 키를 확인하세요.")
        # Gemini는 system/user를 하나의 프롬프트로 합침
        sys_text = "\n\n".join([m["content"] for m in messages if m["role"] == "system"]) or ""
        usr_text = "\n\n".join([m["content"] for m in messages if m["role"] == "user"]) or ""
        prompt = (sys_text + "\n\n" + usr_text).strip()
        try:
            model_obj = self.genai.GenerativeModel(model_name=model)
            resp = model_obj.generate_content(
                prompt,
                generation_config=self.genai.types.GenerationConfig(
                    temperature=temperature,
                    max_output_tokens=max_tokens,
                ),
            )
            if not hasattr(resp, "text") or not resp.text:
                raise RuntimeError("Gemini 응답이 비어 있습니다(안전필터 등).")
            return resp.text.strip()
        except Exception as e:
            raise RuntimeError(f"Gemini 호출 실패: {e}")

# =============================
# Prompts (KR/EN)
# =============================
KR_SYSTEM = (
    "당신은 세계적 기업의 인사담당자 출신 커버레터 코치입니다.\n"
    "- 한국어로 간결하고 명확하게 작성합니다.\n"
    "- 불필요한 과장은 피하고 구체적 성과/지표/역할을 포함합니다.\n"
    "- 문단은 3~5개, 목표 길이 내에서 구성합니다.\n"
    "- 업계 용어는 정확히 사용합니다."
)
EN_SYSTEM = (
    "You are a world-class cover letter coach and former recruiter.\n"
    "- Write concise, professional English.\n"
    "- Avoid fluff; include concrete achievements, metrics, and roles.\n"
    "- Aim for 3–5 paragraphs within target length.\n"
    "- Use industry-appropriate terminology."
)

KR_TONES: Dict[str, str] = {
    "정중한": "톤은 정중하고 신뢰감 있게 유지합니다.",
    "열정적인": "톤은 열정적이고 추진력을 강조합니다.",
    "분석적인": "톤은 논리적이고 분석적으로, 데이터와 근거를 제시합니다.",
    "담백한": "톤은 담백하고 간결하게 핵심만 전달합니다.",
}
EN_TONES: Dict[str, str] = {
    "Polite": "Maintain a polite, trustworthy tone.",
    "Energetic": "Use a passionate, high-drive tone.",
    "Analytical": "Use a logical, analytical tone with data and evidence.",
    "Concise": "Be succinct and straight to the point.",
}

# 메시지 빌더 — 인적사항 포함
def build_draft_messages(lang_code: str, company: str, job_title: str, jd_text: str, resume_text: str,
                         tone_label: str, extra_reqs: str, target_len: int,
                         name: str, age: str, gender: str, address: str) -> List[Dict]:
    personal = f"[인적사항] 이름:{name} / 나이:{age} / 성별:{gender} / 주소:{address}\n".strip()
    if lang_code == "KR":
        system = KR_SYSTEM
        tone_instr = KR_TONES.get(tone_label, "톤은 정중하고 명확합니다.")
        user = (
            f"[회사] {company}\n[직무] {job_title}\n{personal}\n\n"
            f"[채용공고]\n{jd_text}\n\n"
            f"[이력/경험]\n{resume_text}\n\n"
            f"[톤/요청] {tone_instr} {extra_reqs}\n"
            f"[목표 길이] 약 {target_len}자"
        )
    else:
        system = EN_SYSTEM
        tone_instr = EN_TONES.get(tone_label, "Maintain a professional tone.")
        user = (
            f"[Company] {company}\n[Position] {job_title}\n[Personal] Name:{name} / Age:{age} / Gender:{gender} / Address:{address}\n\n"
            f"[Job Description]\n{jd_text}\n\n"
            f"[Resume Highlights]\n{resume_text}\n\n"
            f"[Tone/Guidance] {tone_instr} {extra_reqs}\n"
            f"[Target Length] ~{target_len} words"
        )
    return [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]

def build_refine_messages(lang_code: str, base_text: str, tone: str, target_len: int) -> List[Dict]:
    if lang_code == "KR":
        sys = "너는 전문 에디터이자 HR 컨설턴트다. 중복을 줄이고 구조를 또렷하게 다듬어라."
        user = (
            f"다음 자기소개서를 {tone} 톤에 맞춰 다듬고, 불필요한 군더더기를 제거하고, 문단 구성을 선명히 하세요.\n"
            f"목표 길이는 약 {target_len}자입니다.\n\n[원문]\n{base_text}"
        )
    else:
        sys = "You are an expert editor and HR consultant. Tighten structure, reduce redundancy."
        user = (
            f"Refine the following cover letter into a {tone} tone; reduce fluff and ensure clear paragraphing.\n"
            f"Target length: ~{target_len} words.\n\n[Original]\n{base_text}"
        )
    return [
        {"role": "system", "content": sys},
        {"role": "user", "content": user},
    ]

def build_keywords_messages(lang_code: str, jd_text: str, essay: str) -> List[Dict]:
    if lang_code == "KR":
        sys = "너는 채용담당자다. JD-응답 매칭 품질을 키워드 기준으로 평가한다."
        user = (
            "다음 JD와 자기소개서를 보고, 기업이 중요하게 보는 핵심 키워드 12개 내외를 한국어로 추출하고, "
            "각 키워드의 반영 여부(있음/약함/없음)를 표시해 표 형태의 텍스트로 작성하세요.\n\n"
            f"[JD]\n{jd_text}\n\n[자기소개서]\n{essay}"
        )
    else:
        sys = "You are a recruiter. Evaluate JD-response alignment by keywords."
        user = (
            "From the JD and cover letter, extract ~12 critical keywords and mark for each whether it's Present/Weak/Absent, "
            "as a simple text table.\n\n"
            f"[JD]\n{jd_text}\n\n[Cover Letter]\n{essay}"
        )
    return [
        {"role": "system", "content": sys},
        {"role": "user", "content": user},
    ]

# 영문 변환 메시지
def build_translate_to_en_messages(text_kr: str) -> List[Dict]:
    sys = "You are a professional translator. Translate the Korean cover letter to natural, professional English while preserving structure and intent."
    usr = f"Translate the following Korean cover letter to English:\n\n{text_kr}"
    return [{"role": "system", "content": sys}, {"role": "user", "content": usr}]

# ============== Simple coverage ==============
def naive_coverage(jd_text: str, essay: str) -> Dict[str, float]:
    norm = lambda s: re.sub(r"[^0-9A-Za-z가-힣 ]+", " ", s).lower()
    jd_tokens = [t for t in norm(jd_text).split() if len(t) >= 2]
    es_tokens = [t for t in norm(essay).split() if len(t) >= 2]
    if not jd_tokens or not es_tokens:
        return {"overlap_ratio": 0.0, "jd_vocab": 0, "match_count": 0}
    jd_set, es_set = set(jd_tokens), set(es_tokens)
    match = len(jd_set & es_set)
    return {
        "overlap_ratio": round(match / max(1, len(jd_set)), 3),
        "jd_vocab": len(jd_set),
        "match_count": match,
    }

# =============================
# UI — 전문 서비스 스타일
# =============================
st.set_page_config(page_title="자소서 어렵지 않아용 pro_V3", page_icon="📝", layout="wide")

st.markdown(
    """
    <style>
      .block-container {max-width: 1200px;}
      .app-header {display:flex; align-items:center; justify-content:space-between; padding:8px 4px 16px 4px; border-bottom:1px solid #eee;}
      .title {font-size: 26px; font-weight: 700; letter-spacing: -0.3px;}
      .subtitle {color:#666; font-size:13px;}
      .stTabs [data-baseweb=tab] {height: 48px; padding-top: 10px; font-weight:600;}
      footer {visibility: hidden;}
      .hint {color:#6b7280; font-size:12px;}
    </style>
    """,
    unsafe_allow_html=True,
)

# 상단 헤더 + ⚙️ 설정 팝오버
left, right = st.columns([7, 1])
with left:
    st.markdown('<div class="app-header"><div><div class="title">📝 자소서 어렵지 않아용 pro_V3</div><div class="subtitle">다국어 생성 • 스타일 비교 • 다듬기/분석 • PDF/DOCX</div></div></div>', unsafe_allow_html=True)
with right:
    pop = st.popover("⚙️ 설정", use_container_width=True)
    with pop:
        st.markdown("**모델 및 내보내기 설정**")
        provider_name = st.selectbox("모델 제공자", ["OpenAI", "Gemini"], index=0)
        if provider_name == "OpenAI":
            openai_key = st.text_input("OpenAI API Key", type="password", value=os.getenv("OPENAI_API_KEY", ""))
            model = st.text_input("OpenAI 모델", value="gpt-4o-mini")
            api_key = openai_key
        else:
            gemini_key = st.text_input("Gemini API Key", type="password", value=os.getenv("GEMINI_API_KEY", ""))
            model = st.text_input("Gemini 모델", value="gemini-1.5-flash")
            api_key = gemini_key
        temperature = st.slider("창의성 (Temperature)", 0.0, 1.2, 0.7, 0.05)
        st.markdown("---")
        # 언어/톤/길이도 설정으로 이동
        language = st.selectbox("작성 언어", ["한국어 (KR)", "English (EN)"])
        lang_code = "KR" if language.startswith("한국어") else "EN"
        tone_options = list(KR_TONES.keys()) if lang_code == "KR" else list(EN_TONES.keys())
        default_tones = ["정중한", "분석적인"] if lang_code == "KR" else ["Polite", "Analytical"]
        selected_tones = st.multiselect("비교할 톤 (2–3개 권장)", options=tone_options, default=default_tones)
        target_len = st.slider("목표 길이 (KR:자 / EN:단어)" if lang_code=="KR" else "Target length (words)",
                               min_value=400 if lang_code=="KR" else 150,
                               max_value=1100 if lang_code=="KR" else 450,
                               value=800 if lang_code=="KR" else 280,
                               step=50 if lang_code=="KR" else 10)
        st.markdown("---")
        font_uploader = st.file_uploader("(선택) PDF용 한글 폰트 업로드 (TTF/OTF)", type=["ttf", "otf"])
        font_bytes = font_uploader.read() if font_uploader else None
        font_name = font_uploader.name if font_uploader else None
        st.caption("한국어 PDF가 깨지면 폰트를 업로드하세요 (예: NotoSansKR).")

# 입력 — 인적사항 + 기본 정보 (반투명 placeholder)
colP, colJ = st.columns(2)
with colP:
    st.subheader("인적사항")
    name = st.text_input("이름", placeholder="예) 김준호")
    age = st.text_input("나이", placeholder="예) 30")
    gender = st.text_input("성별", placeholder="예) 남 / 여")
    address = st.text_input("주소", placeholder="예) 서울특별시 강남구 …")
with colJ:
    st.subheader("지원 정보")
    company = st.text_input("회사 / Company", placeholder="예) 카카오엔터프라이즈 (Kakao Enterprise)")
    job_title = st.text_input("직무 / Position", placeholder="예) 데이터 분석가 (Data Analyst)")

jd_text = st.text_area(
    "채용공고 / Job Description",
    height=160,
    placeholder=(
        "예) 주요업무: 데이터 기반 제품 개선, 대시보드 구축, 실험 설계/분석\n"
        "자격요건: SQL 능숙, Python 분석 경험, A/B 테스트 이해, 커뮤니케이션 능력\n"
        "우대사항: 대규모 로그 분석 경험, 클라우드 분석 도구 경험"
    ),
)

resume_text = st.text_area(
    "이력 주요 내용 / Resume Highlights",
    height=160,
    placeholder=(
        "예) 전자상거래 데이터 분석 2년: 전환율 +12% 개선, 리텐션 +8% 향상\n"
        "SQL/BigQuery, Python(pandas, scikit-learn), 실험 설계 경험\n"
        "마케팅/개발과 협업하여 주간 BI 대시보드 운영"
    ),
)

extra_reqs = st.text_area(
    "추가 가이드 (선택)", height=90,
    placeholder="예) 성과는 숫자로, 협업/리더십 1개 문단 강조, 800자 ±15%",
)

# 탭: 1) 생성/내보내기  2) 다듬기/분석
main_tabs = st.tabs(["✍️ 생성 & 내보내기", "🛠️ 다듬기 & 분석"])

# ====== 탭 1: 생성 & 내보내기 ======
with main_tabs[0]:
    st.subheader("스타일 비교 생성")

    provider = OpenAIProvider(api_key) if 'provider_name' in locals() and provider_name == "OpenAI" else GeminiProvider(api_key if 'api_key' in locals() else None)
    disabled = ('api_key' not in locals()) or (not api_key)
    if disabled:
        st.info("🔑 우측 상단 ⚙️에서 API Key와 모델을 설정하세요.")

    go = st.button("변형 생성", disabled=disabled, use_container_width=True)

    gen_results: Dict[str, str] = {}
    if go:
        if not selected_tones:
            st.warning("비교할 톤을 1개 이상 선택하세요.")
        elif not (company and job_title and jd_text and resume_text and name):
            st.warning("회사/직무/JD/이력과 이름을 입력하세요.")
        else:
            with st.spinner("생성 중…"):
                for tone in selected_tones:
                    msgs = build_draft_messages(lang_code, company, job_title, jd_text, resume_text, tone, extra_reqs, target_len, name, age, gender, address)
                    try:
                        txt = provider.generate(msgs, model=model, temperature=temperature, max_tokens=1400)
                    except Exception as e:
                        st.error(str(e))
                        txt = "(생성 실패)"
                    gen_results[tone] = txt or "(생성 실패)"

    if gen_results:
        cols = st.columns(len(gen_results))
        for i, (tone, text) in enumerate(gen_results.items()):
            with cols[i]:
                st.markdown(f"**{tone}**")
                st.text_area("", value=text, height=360, key=f"out_{tone}")

        st.markdown("---")
        st.subheader("내보내기 / 영문 변환")
        exp_cols = st.columns(len(gen_results))
        for i, (tone, text) in enumerate(gen_results.items()):
            with exp_cols[i]:
                label = f"{company} — {job_title} ({tone})"
                # PDF
                try:
                    pdf_bytes = build_pdf_bytes(text, label, "KR" if lang_code=="KR" else "EN", font_bytes if 'font_bytes' in locals() else None, font_name if 'font_name' in locals() else None)
                    st.download_button("⬇️ PDF", data=pdf_bytes, file_name=f"{company}_{job_title}_{tone}.pdf")
                except Exception as e:
                    st.error(f"PDF 실패: {e}")
                # DOCX
                if HAS_DOCX:
                    try:
                        buf = io.BytesIO()
                        doc = Document()
                        doc.add_heading(label, level=1)
                        for para in text.split("\n\n"):
                            doc.add_paragraph(para)
                        doc.save(buf)
                        buf.seek(0)
                        st.download_button("⬇️ DOCX", data=buf.getvalue(), file_name=f"{company}_{job_title}_{tone}.docx")
                    except Exception as e:
                        st.error(f"DOCX 실패: {e}")
                else:
                    st.caption("DOCX: `pip install python-docx` 필요")

        # 통합 PDF
        if st.checkbox("모든 버전을 하나의 PDF로 묶기"):
            all_text = "\n\n\n".join([f"[{tone}]\n\n{text}" for tone, text in gen_results.items()])
            try:
                combined = build_pdf_bytes(all_text, f"{company} — {job_title} (All styles)", "KR" if lang_code=="KR" else "EN", font_bytes if 'font_bytes' in locals() else None, font_name if 'font_name' in locals() else None)
                st.download_button("⬇️ 통합 PDF", data=combined, file_name=f"{company}_{job_title}_ALL.pdf")
            except Exception as e:
                st.error(f"통합 PDF 실패: {e}")

        # 영문 변환(각 버전)
        st.markdown("---")
        st.subheader("생성본 영문 변환")
        trans_cols = st.columns(len(gen_results))
        for i, (tone, text) in enumerate(gen_results.items()):
            with trans_cols[i]:
                if st.button(f"영문으로 변환 ({tone})", key=f"tr_btn_{tone}", use_container_width=True, disabled=disabled):
                    provider_tr = OpenAIProvider(api_key) if provider_name == "OpenAI" else GeminiProvider(api_key)
                    with st.spinner("영문 변환 중…"):
                        try:
                            msgs_tr = build_translate_to_en_messages(text)
                            en_text = provider_tr.generate(msgs_tr, model=model, temperature=0.3, max_tokens=1400)
                            st.text_area("영문 결과", value=en_text, height=300, key=f"out_en_{tone}")
                        except Exception as e:
                            st.error(str(e))

# ====== 탭 2: 다듬기 & 분석 ======
with main_tabs[1]:
    st.subheader("다듬기 / 재톤 적용")
    st.markdown("<span class='hint'>예시: 아래에 기존 자기소개서 초안을 붙여넣고 톤과 길이를 조절해 다듬을 수 있어요.</span>", unsafe_allow_html=True)

    base_text = st.text_area("원문 붙여넣기", height=220, placeholder="예) 저는 전자상거래 데이터 분석가로서 2년간 전환율 개선과 리텐션 향상을 이끌었습니다…")

    colR1, colR2, colR3 = st.columns([2, 1, 1])
    with colR1:
        refine_tone_opts = (tone_options)
        refine_tone = st.selectbox("톤 선택", options=refine_tone_opts, index=0)
    with colR2:
        ref_len = st.number_input("목표 길이", value=800 if 'target_len' in locals() and lang_code=="KR" else 280, min_value=200, max_value=2000, step=50)
    with colR3:
        btn_refine = st.button("다듬기 실행", disabled=('api_key' not in locals()) or (not api_key), use_container_width=True)

    refined_text = ""
    if btn_refine:
        if not base_text.strip():
            st.warning("원문을 입력하세요.")
        else:
            provider2 = OpenAIProvider(api_key) if provider_name == "OpenAI" else GeminiProvider(api_key)
            msgs = build_refine_messages(lang_code, base_text, refine_tone, int(ref_len))
            with st.spinner("다듬는 중…"):
                try:
                    refined_text = provider2.generate(msgs, model=model, temperature=0.6, max_tokens=1200)
                except Exception as e:
                    st.error(str(e))
                    refined_text = "(다듬기 실패)"
    st.text_area("다듬은 결과", value=refined_text, height=220, key="refined_view")

    st.markdown("---")
    st.subheader("분석: 키워드 & 겹침")
    st.markdown("<span class='hint'>예시: JD와 다듬은 결과를 기준으로 핵심 키워드와 겹침 비율을 확인합니다.</span>", unsafe_allow_html=True)

    colK1, colK2 = st.columns(2)
    with colK1:
        analysis_src = st.text_area("분석 대상 본문", value=refined_text, height=180, placeholder="예) 다듬은 결과를 여기에 붙여넣으세요")
    with colK2:
        btn_kw = st.button("키워드 추출 및 겹침 계산", disabled=('api_key' not in locals()) or (not api_key), use_container_width=True)
        analysis_out = st.empty()

    if btn_kw:
        if not (jd_text and (analysis_src or refined_text)):
            analysis_out.warning("JD와 분석 본문을 확인하세요.")
        else:
            provider3 = OpenAIProvider(api_key) if provider_name == "OpenAI" else GeminiProvider(api_key)
            msgs = build_keywords_messages(lang_code, jd_text, analysis_src or refined_text)
            with st.spinner("분석 중…"):
                # LLM keywords
                kw_text = ""
                try:
                    kw_text = provider3.generate(msgs, model=model, temperature=0.2, max_tokens=800)
                except Exception as e:
                    analysis_out.error(f"키워드 추출 실패: {e}")
                # Naive coverage
                try:
                    cov = naive_coverage(jd_text, analysis_src or refined_text)
                    st.metric("JD 어휘 겹침 비율", f"{cov['overlap_ratio']*100:.1f}%")
                    st.caption(f"JD 어휘수: {cov['jd_vocab']} / 매치: {cov['match_count']}")
                except Exception as e:
                    st.error(f"겹침 계산 실패: {e}")
            st.text_area("키워드 평가 (LLM)", value=kw_text, height=180)

st.divider()
st.markdown("<small>Tip: 우측 상단 ⚙️에서 언어/톤/길이까지 한 번에 설정하고, 탭별 예시를 참고해 바로 생성해 보세요.</small>", unsafe_allow_html=True)
